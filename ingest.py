import os
import json
import csv
import pdfplumber
from docx import Document as DocxDocument
from typing import List
from langchain.vectorstores import Chroma
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document

# Constants
PERSIST_DIRECTORY = "/tmp/chroma_db"
VALID_EXTENSIONS = [".pdf", ".docx", ".json", ".csv"]

# Components
text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

# ---- Local Parsers ----

def parse_pdf(file_path: str) -> List[Document]:
    texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
    return [Document(page_content=page) for page in texts]

def parse_docx(file_path: str) -> List[Document]:
    doc = DocxDocument(file_path)
    full_text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return [Document(page_content=full_text)]

def parse_json(file_path: str) -> List[Document]:
    with open(file_path, 'r', encoding='utf-8') as f:
        data = json.load(f)

    docs = []
    if isinstance(data, list):
        for entry in data:
            content = json.dumps(entry, indent=2)
            docs.append(Document(page_content=content))
    elif isinstance(data, dict):
        docs.append(Document(page_content=json.dumps(data, indent=2)))
    else:
        raise ValueError("Unsupported JSON structure.")
    
    return docs

def parse_csv(file_path: str) -> List[Document]:
    docs = []
    with open(file_path, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            content = json.dumps(row, indent=2)
            docs.append(Document(page_content=content))
    return docs

# ---- Ingestion Logic ----

def ingest_file(file_path: str) -> List[Document]:
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        print(f"📄 Parsing PDF locally...")
        return parse_pdf(file_path)
    elif ext == ".docx":
        print(f"📄 Parsing DOCX locally...")
        return parse_docx(file_path)
    elif ext == ".json":
        print("🧾 Parsing JSON...")
        return parse_json(file_path)
    elif ext == ".csv":
        print("📊 Parsing CSV...")
        return parse_csv(file_path)
    else:
        print(f"❌ Skipping unsupported file: {file_path}")
        return []

def ingest_folder(folder_path: str):
    all_docs = []

    print(f"📁 Ingesting files from: {folder_path}")
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if not os.path.isfile(file_path):
            continue
        if os.path.splitext(file_path)[1].lower() not in VALID_EXTENSIONS:
            print(f"⏭️ Skipping {filename} (unsupported extension)")
            continue
        
        docs = ingest_file(file_path)
        all_docs.extend(docs)

    print(f"🔍 Splitting all documents into chunks...")
    chunks = text_splitter.split_documents(all_docs)

    print(f"💾 Storing {len(chunks)} chunks in ChromaDB...")
    vectordb = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=PERSIST_DIRECTORY
    )
    vectordb.persist()
    print("✅ All documents ingested and saved!")

# Optional: direct run
if __name__ == "__main__":
    ingest_folder("documents")
